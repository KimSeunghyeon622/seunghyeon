from docx import Document

doc = Document('code_review_result.docx')

with open('code_review_result.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + '\n')

    # 테이블도 확인
    for table in doc.tables:
        f.write('\n--- TABLE ---\n')
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            f.write(' | '.join(row_text) + '\n')

print("Done - saved to code_review_result.txt")
